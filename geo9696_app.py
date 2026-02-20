import streamlit as st
import fitz  # PyMuPDF
from docx import Document
import io
import os
import base64
import pandas as pd
from datetime import datetime

# --- 1. CONFIGURATION & DIRECTORIES ---
# Setting up the workspace for Geography 9696
SUBJECT_CODE = "9696"
SAVE_DIR = "pyp9696_qp"
MS_DIR = "pyp9696_ms"
DIAGRAM_DIR = "geography_diagrams"
GALLERY_FILE = "geography_case_studies.csv"
GLOSSARY_FILE = "geography_glossary.csv"

# Ensure all local storage folders exist
for folder in [SAVE_DIR, MS_DIR, DIAGRAM_DIR]:
    if not os.path.exists(folder):
        os.makedirs(folder)

# Syllabus structure for 9696
geo_topics = {
    "AS Physical Core": ["Hydrology", "Fluvial geomorphology", "Atmosphere", "Weather", "Rocks", "Weathering"],
    "AS Human Core": ["Population", "Migration", "Settlement dynamics"],
    "A2 Physical Options": ["Tropical environments", "Coastal environments", "Hazardous environments", "Hot arid",
                            "Semi-arid"],
    "A2 Human Options": ["Production", "Environmental management", "Global interdependence", "Economic transition"]
}

variants_map = {
    "1": ["11", "12", "13"], "2": ["21", "22", "23"],
    "3": ["31", "32", "33"], "4": ["41", "42", "43"]
}

# --- 2. CORE ENGINE FUNCTIONS ---
def extract_questions(path, keyword):
    """Searches PDF for keywords and returns text by page."""
    if not os.path.exists(path): return None
    try:
        doc = fitz.open(path)
        output = ""
        for page in doc:
            text = page.get_text()
            if not keyword or keyword.lower() in text.lower():
                output += f"\n--- {os.path.basename(path)} (P.{page.number + 1}) ---\n{text}"
        doc.close()
        return output if output.strip() else None
    except Exception as e:
        return f"Error reading PDF: {e}"


def display_pdf(file_path):
    """Renders PDF in Streamlit iframe."""
    with open(file_path, "rb") as f:
        base64_pdf = base64.b64encode(f.read()).decode('utf-8')
    pdf_display = f'<iframe src="data:application/pdf;base64,{base64_pdf}" width="100%" height="800" type="application/pdf"></iframe>'
    st.markdown(pdf_display, unsafe_allow_html=True)

##############################################################
def save_to_gallery(topic, content, metadata):
    try:
        new_entry = pd.DataFrame([{
            "Date Saved": datetime.now().strftime("%Y-%m-%d"),
            "Topic": topic,
            "Source": metadata,
            "Content": content
        }])
        # This ensures the file is saved exactly where the script is
        file_path = os.path.join(os.getcwd(), GALLERY_FILE)

        if os.path.exists(file_path):
            df = pd.read_csv(file_path)
            df = pd.concat([df, new_entry], ignore_index=True)
        else:
            df = new_entry

        df.to_csv(file_path, index=False)
        st.success(f"✅ Saved to: {file_path}")  # This tells us exactly where it went!
    except Exception as e:
        st.error(f"❌ Save failed: {e}")

def save_to_glossary(term, definition):
    """Saves definitions to CSV."""
    new_entry = pd.DataFrame([{"Term": term, "Definition": definition}])
    df = pd.concat([pd.read_csv(GLOSSARY_FILE), new_entry] if os.path.exists(GLOSSARY_FILE) else [new_entry],
                   ignore_index=True).drop_duplicates(subset=['Term'])
    df.to_csv(GLOSSARY_FILE, index=False)

def analyze_predictions(current_year):
    """Calculates topic priority based on last seen date in Gallery."""
    if not os.path.exists(GALLERY_FILE): return None
    df = pd.read_csv(GALLERY_FILE)
    df['Year'] = df['Source'].str.extract(r'(\d{4})').astype(float)
    analysis = []
    for cat, units in geo_topics.items():
        for unit in units:
            last_seen = df[df['Topic'].str.contains(unit, case=False)]['Year'].max()
            if pd.isna(last_seen):
                status = "No Data"; color = "gray"
            elif (current_year - last_seen) >= 2:
                status = "⚠️ High Priority"; color = "red"
            elif (current_year - last_seen) == 1:
                status = "🟡 Medium Priority"; color = "orange"
            else:
                status = "🟢 Low Priority"; color = "green"
            analysis.append({"Component": cat, "Unit": unit, "Last Examined": last_seen, "Priority": status})
    return pd.DataFrame(analysis)

# --- 3. STREAMLIT INTERFACE ---
st.set_page_config(page_title="Geography 9696 Tutor Portal", layout="wide")
st.title("🌍 Geography 9696 Official Tutor Resource Platform")

# SIDEBAR: MANAGEMENT
with st.sidebar:
    st.header("📤 Resource Management")
    u_y = st.number_input("Exam Year", 2018, 2030, 2026)
    u_s = st.selectbox("Exam Session", ["MARCH (m)", "JUNE (s)", "NOVEMBER (w)"])
    u_p = st.selectbox("Paper Number", ["1", "2", "3", "4"])
    u_v = st.selectbox("Variant", variants_map[u_p])
    u_file = st.file_uploader("Upload PDF File", type="pdf")
    u_type = st.radio("Upload Category", ["Question Paper (QP)", "Marking Scheme (MS)"])

    if st.button("Add to Database"):
        if u_file:
            s_let = u_s.split('(')[1][0]
            prefix = "qp" if u_type == "Question Paper (QP)" else "ms"
            new_fn = f"{SUBJECT_CODE}_{s_let}{str(u_y)[-2:]}_{prefix}_{u_v}.pdf"
            path = os.path.join(SAVE_DIR if prefix == "qp" else MS_DIR, new_fn)
            with open(path, "wb") as f:
                f.write(u_file.getbuffer())
            st.success(f"Successfully Stored: {new_fn}")
        else:
            st.error("Please upload a file first.")

# MAIN TABS
t1, t2, t3, t4, t5, t6 = st.tabs([
    "🔍 Search", "📚 Batch", "🖼️ Case Studies", "📝 Revision",
    "📊 Diagrams", "📈 Predictor"
])

# TAB 1: SEARCH & CLIP
with t1:
    st.subheader("🔍 Find Exam Content")
    col1, col2 = st.columns([2, 1])
    with col1:
        cat = st.selectbox("Syllabus Component", list(geo_topics.keys()))
        unit = st.selectbox("Core Unit", geo_topics[cat])
        s_topic = st.text_input("Refine Keyword Search", value=unit)
    with col2:
        s_yr = st.selectbox("Year", range(2018, 2027), index=8)
        s_p = st.selectbox("Select Paper", ["1", "2", "3", "4"])
        s_v = st.selectbox("Select Variant", variants_map[s_p])

    # TRIGGER SEARCH
    if st.button("Search Papers"):
        found_data = []
        for s in ["m", "s", "w"]:
            fn = f"{SUBJECT_CODE}_{s}{str(s_yr)[-2:]}_qp_{s_v}.pdf"
            path = os.path.join(SAVE_DIR, fn)
            res = extract_questions(path, s_topic)
            if res:
                found_data.append({"text": res, "src": f"{s_yr} P{s_p} V{s_v}", "session": s})

        # LOCK results into memory
        st.session_state['geo_search_results'] = found_data

    # DISPLAY results from memory
    if 'geo_search_results' in st.session_state:
        for i, item in enumerate(st.session_state['geo_search_results']):
            st.info(item['text'][:600] + "...")

            # Two columns for our two buttons
            c1, c2 = st.columns(2)

            with c1:
                if st.button(f"📌 Save Snippet {i + 1}", key=f"save_btn_{i}"):
                    save_to_gallery(s_topic, item['text'], item['src'])
                    st.toast(f"Saved {s_topic}!")

            with c2:
                # NEW: Tiny button for Answer Scheme
                # Matches your filename: 9696_s25_ms_11.pdf
                year_short = str(item['src'].split()[0])[-2:]
                ms_fn = f"{SUBJECT_CODE}_{item['session']}{year_short}_ms_{s_v}.pdf"
                ms_path = os.path.join(MS_DIR, ms_fn)

                if os.path.exists(ms_path):
                    with open(ms_path, "rb") as f:
                        st.download_button("📂 View Answer Scheme", f, file_name=ms_fn, key=f"ms_btn_{i}")
                else:
                    st.caption("MS not found in pyp9696_ms")

# TAB 2: BATCH EXTRACTION
with t2:
    st.subheader("📚 Create Topical Booklets")
    b_start = st.number_input("Batch Start Year", 2018, 2025, 2022)
    b_topic = st.text_input("Booklet Topic (e.g. 'Hazards')")
    if st.button("🚀 Compile 3-Year Booklet"):
        all_text = ""
        for yr in range(b_start, b_start + 4):
            for s in ["m", "s", "w"]:
                # Logic iterates through all variants for that paper type
                for v in variants_map["1"]:  # Checks common variants
                    fn = f"{SUBJECT_CODE}_{s}{str(yr)[-2:]}_qp_{v}.pdf"
                    path = os.path.join(SAVE_DIR, fn)
                    res = extract_questions(path, b_topic)
                    if res: all_text += f"\n\n{'=' * 40}\nYEAR: {yr} | SESSION: {s.upper()}\n{'=' * 40}\n{res}"

        if all_text:
            doc = Document();
            doc.add_heading(f'9696 Geography: {b_topic}', 0);
            doc.add_paragraph(all_text)
            bio = io.BytesIO();
            doc.save(bio)
            st.download_button("📥 Download Word Booklet", bio.getvalue(), f"9696_{b_topic}_Booklet.docx")
        else:
            st.warning("No questions found for this range.")

# TAB 3: CASE STUDY BANK
with t3:
    st.subheader("🖼️ Case Study Bank")

    if os.path.exists(GALLERY_FILE):
        df = pd.read_csv(GALLERY_FILE)

        if not df.empty:
            st.write(f"✅ Found {len(df)} entries")

            # We create a loop to display each entry with a delete option
            for i, row in df.iterrows():
                col1, col2 = st.columns([0.8, 0.2])

                with col1:
                    with st.expander(f"📌 {row['Topic']} ({row['Date Saved']})"):
                        st.write(f"**Source:** {row['Source']}")
                        st.write(row['Content'])

                with col2:
                    # Unique key for each delete button based on index
                    if st.button("🗑️ Delete", key=f"del_{i}"):
                        df = df.drop(i)
                        df.to_csv(GALLERY_FILE, index=False)
                        st.rerun()
        else:
            st.warning("Gallery is currently empty.")
    else:
        st.info("No database file found yet.")

###########################################################
# --- TAB 4: REVISION GENERATOR (SYNCED + PREVIEW) ---
# --- TAB 4: REVISION GENERATOR (FORCED UI UPDATE) ---
with t4:
    st.subheader("📝 Handout Creator")
    if os.path.exists(GALLERY_FILE):
        gal_df = pd.read_csv(GALLERY_FILE)

        if not gal_df.empty:
            # 1. Choose the Case Study text
            choice = st.selectbox("1. Select Evidence", gal_df.index,
                                  format_func=lambda x: f"{gal_df.iloc[x]['Topic']} ({gal_df.iloc[x]['Source']})")

            # 2. THIS IS THE NEW LINE YOU NEED TO SEE ON SCREEN
            diag_files = [f for f in os.listdir(DIAGRAM_DIR) if f.endswith(('.png', '.jpg', '.jpeg'))]
            selected_diag = st.selectbox("2. Choose a Diagram from Library", ["None"] + diag_files)

            # 3. The Button
            if st.button("🔨 Generate Handout"):
                selected = gal_df.iloc[choice]
                doc = Document()
                doc.add_heading(f"9696 Revision: {selected['Topic']}", 0)
                doc.add_paragraph(selected['Content'])

                # Image Insertion Logic
                if selected_diag != "None":
                    doc.add_heading("Refer to the Diagram Below", level=1)
                    diag_path = os.path.join(DIAGRAM_DIR, selected_diag)
                    from docx.shared import Inches

                    doc.add_picture(diag_path, width=Inches(4))

                # Tasks
                doc.add_heading("Tasks", level=1)
                doc.add_paragraph("1. Identify the processes shown in the diagram.")
                doc.add_paragraph("2. Evaluate how this affects the 9696 Case Study.")

                bio = io.BytesIO()
                doc.save(bio)
                st.download_button("📥 Download Handout", bio.getvalue(), "Revision_Sheet.docx")
    else:
        st.warning("Please save some snippets first!")
##########################################################

# TAB 5: DIAGRAM LIBRARY
with t5:
    st.subheader("📊 Diagram Library")
    diag_up = st.file_uploader("Upload Diagram Image", type=['png', 'jpg', 'jpeg'])
    diag_name = st.text_input("Diagram Label")
    if st.button("Upload to Library"):
        if diag_up and diag_name:
            # Clean the name to avoid file errors
            clean_name = diag_name.replace(' ', '_')
            path = os.path.join(DIAGRAM_DIR, f"{clean_name}.png")
            with open(path, "wb") as f:
                f.write(diag_up.getbuffer())
            st.success("Diagram Added!")
            st.rerun()

    st.divider()

    # Display images with a delete button
    files = os.listdir(DIAGRAM_DIR)
    if files:
        cols = st.columns(3)
        for i, f_name in enumerate(files):
            with cols[i % 3]:
                img_path = os.path.join(DIAGRAM_DIR, f_name)
                st.image(img_path)

                # Column split for label and delete button
                c_lab, c_del = st.columns([0.8, 0.2])
                with c_lab:
                    st.caption(f_name.replace("_", " ").split(".")[0])
                with c_del:
                    # Logic to delete the file from your computer
                    if st.button("🗑️", key=f"del_diag_{i}"):
                        os.remove(img_path)
                        st.toast(f"Deleted {f_name}")
                        st.rerun()
    else:
        st.info("No diagrams uploaded yet.")

with t6:
    st.subheader("📈 Exam Question Predictor")
    pred_df = analyze_predictions(2026)
    if pred_df is not None:
        def style_priority(v):
            if "High" in v: return 'color: red; font-weight: bold'
            if "Medium" in v: return 'color: orange'
            return 'color: green'

        st.dataframe(pred_df.style.applymap(style_priority, subset=['Priority']))
        st.info("Priority is based on how long it has been since a topic was last saved to your Case Study Bank.")
    else:
        st.warning("Save snippets to the Case Study Bank to enable prediction analysis.")
